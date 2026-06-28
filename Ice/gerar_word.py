import os
import pathlib
import docx
import shutil
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

WORKSPACE = r"c:\Users\3D_OCT\Documents\Antigravity\Ice"
DEST_DIR = r"D:\ICE_Apresentacao"

pathlib.Path(WORKSPACE).mkdir(parents=True, exist_ok=True)
pathlib.Path(DEST_DIR).mkdir(parents=True, exist_ok=True)

DOC_PATH_LOCAL = os.path.join(WORKSPACE, "ICE_Projeto_Pesquisa_Ambrosio.docx")
DOC_PATH_DEST = os.path.join(DEST_DIR, "ICE_Projeto_Pesquisa_Ambrosio.docx")

# Colors
COLOR_NAVY = RGBColor(11, 26, 46)     # #0b1a2e
COLOR_TEAL = RGBColor(23, 161, 165)    # #17a1a5
COLOR_GOLD = RGBColor(201, 168, 76)    # #c9a84c
COLOR_GRAY = RGBColor(74, 95, 115)     # #4a5f73
COLOR_BLACK = RGBColor(30, 45, 66)     # #1e2d42

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_left_border(cell, hex_color, sz="36"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{hex_color}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)

def add_callout(doc, text, cite):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.cell(0, 0)
    
    set_cell_background(cell, "FDF9F0")
    set_cell_left_border(cell, "C9A84C", "36")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=140)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = COLOR_BLACK
    
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(cite)
    run2.font.name = 'Arial'
    run2.font.size = Pt(9)
    run2.font.color.rgb = COLOR_GRAY

def build_document():
    doc = docx.Document()
    
    # Configure page setup
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Title Block (Simulating header banner)
    p_tag = doc.add_paragraph()
    p_tag.paragraph_format.space_before = Pt(0)
    p_tag.paragraph_format.space_after = Pt(6)
    run_tag = p_tag.add_run("PROJETO DE PESQUISA · BRAIN / RIO DE JANEIRO · JUNHO 2026")
    run_tag.font.name = 'Arial'
    run_tag.font.size = Pt(8.5)
    run_tag.bold = True
    run_tag.font.color.rgb = COLOR_TEAL
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(8)
    p_title.paragraph_format.line_spacing = 1.15
    run_title = p_title.add_run("Índice de Coerência do Eixo (ICE):\nUm preditor pré-operatório de resultado visual no implante de anel intraestromal")
    run_title.font.name = 'Georgia'
    run_title.font.size = Pt(20)
    run_title.bold = True
    run_title.font.color.rgb = COLOR_NAVY
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(20)
    run_sub = p_sub.add_run("Documento científico detalhado para avaliação de relevância clínica e validação do protocolo retrospectivo/prospectivo em ceratocone.")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(10.5)
    run_sub.italic = True
    run_sub.font.color.rgb = COLOR_GRAY
    
    # Meta table
    tbl_meta = doc.add_table(rows=2, cols=2)
    tbl_meta.autofit = True
    
    meta_data = [
        ("Investigador Principal:", "Miguel Reis — Biomecânica Corneana Computacional", 
         "Destinado para Parecer de:", "Prof. Dr. Renato Ambrósio Jr. — BrAIn"),
        ("Data de Submissão:", "Junho 2026", 
         "Objetivo do Envio:", "Apresentação e Parecer de Relevância Científica")
    ]
    
    for r_idx, row_data in enumerate(meta_data):
        c0_title, c0_val, c1_title, c1_val = row_data
        
        # Col 0
        cell0 = tbl_meta.cell(r_idx, 0)
        p0 = cell0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(4)
        run_t0 = p0.add_run(c0_title + " ")
        run_t0.bold = True
        run_t0.font.size = Pt(9.5)
        run_t0.font.color.rgb = COLOR_NAVY
        run_v0 = p0.add_run(c0_val)
        run_v0.font.size = Pt(9.5)
        run_v0.font.color.rgb = COLOR_BLACK
        
        # Col 1
        cell1 = tbl_meta.cell(r_idx, 1)
        p1 = cell1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(4)
        run_t1 = p1.add_run(c1_title + " ")
        run_t1.bold = True
        run_t1.font.size = Pt(9.5)
        run_t1.font.color.rgb = COLOR_NAVY
        run_v1 = p1.add_run(c1_val)
        run_v1.font.size = Pt(9.5)
        run_v1.font.color.rgb = COLOR_BLACK
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ------------------
    # Helper to add headings
    # ------------------
    def add_sec_heading(num, title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        
        run_num = p.add_run(f"SEÇÃO {num}: ")
        run_num.font.name = 'Arial'
        run_num.font.size = Pt(9)
        run_num.bold = True
        run_num.font.color.rgb = COLOR_TEAL
        
        run_title = p.add_run(title)
        run_title.font.name = 'Georgia'
        run_title.font.size = Pt(14)
        run_title.bold = True
        run_title.font.color.rgb = COLOR_NAVY
        
        # Add a bottom border line in word
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="4" w:color="E2EAF4"/></w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)

    def add_p(text, bold_prefix=None, space_after=8, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_pre = p.add_run(bold_prefix)
            run_pre.bold = True
            run_pre.font.name = 'Arial'
            run_pre.font.size = Pt(10)
            run_pre.font.color.rgb = COLOR_NAVY
            
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.italic = italic
        run.font.color.rgb = COLOR_BLACK
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_pre = p.add_run(bold_prefix)
            run_pre.bold = True
            run_pre.font.name = 'Arial'
            run_pre.font.size = Pt(10)
            run_pre.font.color.rgb = COLOR_NAVY
            
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_BLACK
        return p

    # 1. SUMÁRIO EXECUTIVO
    add_sec_heading("1", "Sumário Executivo")
    add_callout(doc, 
                '"O Kmax informa a altura do incêndio. O ICE informa se a saída de emergência ainda está alinhada."', 
                "Analogia central do projeto: gravidade estrutural e coerência funcional são dimensões independentes.")
    
    add_p("O Índice de Coerência do Eixo (ICE) é um índice pré-operatório desenvolvido para responder a uma pergunta que os índices existentes não respondem: se o anel intraestromal corrigir a curvatura, o sistema visual deste paciente conseguirá converter essa melhora geométrica em ganho visual real?")
    add_p("O paradoxo que motivou o projeto é documentado clinicamente: pacientes com Kmax idêntico, mesma paquimetria, mesmo estadiamento — e respostas visuais radicalmente diferentes ao anel. A hipótese do ICE é que a causa dessa variabilidade está no desacoplamento entre três eixos ópticos: o eixo de maior curvatura (topografia), o eixo do cilindro manifesto (refração) e o eixo da aberração comática (aberrometria). Quando esses três eixos apontam em direções diferentes, a melhora geométrica produzida pelo anel se fragmenta ao longo do sistema óptico e não se integra em ganho de acuidade.")
    
    add_p("Indicadores Base do Projeto:", space_after=4)
    add_bullet(" pares clínicos pareados com Pentacam + OCT Triton.", "443")
    add_bullet(" perfis de curvatura analisados via elevação por polinômios ortogonais (SPR).", "660")
    add_bullet(" modelos de elementos finitos (FEM) com material anisotrópico HGO.", "24")
    add_bullet(" de Área Sob a Curva ROC retrospectiva na predição de ganho >= 3 linhas de visão.", "0,82")

    # 2. PROBLEMA E HIPÓTESE
    add_sec_heading("2", "Problema e Hipótese")
    add_p("A literatura de ICRS em ceratocone documenta consistentemente uma variabilidade clínica que os nomogramas atuais não explicam. Estudos com mais de 300 olhos mostram que, entre pacientes operados com indicação tecnicamente adequada, cerca de 27% obtêm ganho de BCVA inferior a duas linhas, mesmo com melhora objetiva da topografia. Essa dissociação entre resultado topográfico e resultado visual é o problema central que o ICE busca predizer.")
    
    # Hypothesis Box (Simulated with 1-cell table)
    tbl_hyp = doc.add_table(rows=1, cols=1)
    tbl_hyp.columns[0].width = Inches(6.5)
    cell_hyp = tbl_hyp.cell(0,0)
    set_cell_background(cell_hyp, "F0F7FF")
    set_cell_left_border(cell_hyp, "2892D7", "24")
    set_cell_margins(cell_hyp, top=100, bottom=100, left=150, right=150)
    p_hyp = cell_hyp.paragraphs[0]
    run_hl = p_hyp.add_run("HIPÓTESE PRINCIPAL:\n")
    run_hl.bold = True
    run_hl.font.size = Pt(9)
    run_hl.font.color.rgb = COLOR_TEAL
    run_ht = p_hyp.add_run("O grau de coerência entre o eixo topográfico, o eixo refrativo e o eixo comático — quantificado pelo ICE — é um preditor independente do ganho de BCVA após ICRS, com desempenho discriminativo superior ao Kmax isolado (AUC > 0,75 no estudo prospectivo).")
    run_ht.font.size = Pt(10)
    run_ht.font.color.rgb = COLOR_NAVY
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    add_p("O Kmax mede a magnitude da deformação corneana, mas não mede como essa deformação se propaga pelo sistema óptico. Dois olhos com Kmax de 62 D podem ter:", bold_prefix="Por que o Kmax é insuficiente: ")
    add_bullet(" e cilindro manifesto coincidentes (ICE alto) — sistema coerente, ganho esperado.", "Eixo topográfico")
    add_bullet(", cilindro e coma em meridianos diferentes (ICE baixo) — sistema fragmentado, ganho improvável.", "Eixo topográfico")
    
    add_p("Na análise de 660 perfis SPR, a correlação espacial entre os ápices da superfície anterior e posterior da córnea foi de r = 0,028 (p = 0,50) — estatisticamente nula. O ápice posterior está, em média, 2,6 mm mais periférico, e em posição angular diferente. Isso significa que o eixo de referência convencional (K-steep) pode estar sistematicamente deslocado em relação ao eixo real do cone.", bold_prefix="O achado que sustenta a hipótese: ")

    # 3. O ÍNDICE ICE — DEFINIÇÃO FORMAL
    add_sec_heading("3", "O Índice ICE — Definição Formal")
    add_p("O ICE é um índice contínuo entre 0 e 1, produto de três componentes independentes, cada um capturando uma dimensão distinta da coerência óptico-geométrica.")
    
    # Formula Box
    tbl_form = doc.add_table(rows=1, cols=1)
    tbl_form.columns[0].width = Inches(6.5)
    cell_form = tbl_form.cell(0,0)
    set_cell_background(cell_form, "0B1A2E")
    set_cell_margins(cell_form, top=140, bottom=140, left=180, right=180)
    
    pf1 = cell_form.paragraphs[0]
    run_ft = pf1.add_run("FÓRMULA ICE 2.0 — TRIÁDICO\n\n")
    run_ft.bold = True
    run_ft.font.size = Pt(9)
    run_ft.font.color.rgb = COLOR_TEAL
    
    run_fe = pf1.add_run("ICE = ICE_astig  x  Optical_coherence  x  ICE_axis_triad\n")
    run_fe.bold = True
    run_fe.font.name = 'Georgia'
    run_fe.font.size = Pt(12)
    run_fe.font.color.rgb = RGBColor(255, 255, 255)
    
    run_fe2 = pf1.add_run("      = [1 - |Astig_T - Cyl| / (Astig_T + Cyl)]  x  [LOA / (LOA + HOA)]  x  [1 - D_max / 90°]\n\n")
    run_fe2.font.name = 'Georgia'
    run_fe2.font.size = Pt(10.5)
    run_fe2.font.color.rgb = RGBColor(220, 232, 245)
    
    run_fn = pf1.add_run("Onde:\n• D_max = max( |θ_topo - θ_cyl|, |θ_topo - θ_coma|, |θ_cyl - θ_coma| )\n• Astig_T = astigmatismo topográfico  ·  Cyl = cilindro manifesto\n• LOA = aberrações de baixa ordem  ·  HOA = aberrações de alta ordem")
    run_fn.font.size = Pt(8.5)
    run_fn.font.color.rgb = RGBColor(180, 200, 220)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    add_p("Compara a magnitude do astigmatismo topográfico com o cilindro manifesto. Dissociação entre os dois indica que o sistema refrativo não está 'lendo' o que a córnea impõe — frequentemente por adaptação neural ou erro de medida.", bold_prefix="Bloco 1 — Coerência de magnitude: ")
    add_p("Razão entre aberrações de baixa ordem (corrigíveis) e totais. Quando aberrações de alta ordem dominam, halos e imagens duplas persistem mesmo com boa refração manifesta — o ganho de BCVA é estruturalmente limitado.", bold_prefix="Bloco 2 — Pureza óptica: ")
    add_p("Maior discordância angular entre os três eixos (topográfico, refrativo, comático). Eixos a <15°: alta coerência. A >30°: o anel posicionado no K-steep pode aumentar o coma ao invés de reduzi-lo.", bold_prefix="Bloco 3 — Coerência angular triádica: ")
    
    # Classification Table
    add_p("Classificação Clínica e Condutas do ICE:", space_after=4)
    tbl_class = doc.add_table(rows=4, cols=4)
    tbl_class.style = 'Light Shading Accent 1'
    
    headers = ["Faixa ICE", "Classificação", "Interpretação", "Conduta Proposta"]
    for i, h in enumerate(headers):
        tbl_class.cell(0, i).paragraphs[0].add_run(h).bold = True
        
    class_rows = [
        (">= 0,65", "Tipo 1 — Coerente", "Eixos alinhados, HOA controladas, magnitude concordante", "Operar. Maximizar VEsférico. Prometer ganho objetivo."),
        ("0,50 – 0,64", "Zona cinzenta", "Coerência parcial; pelo menos um bloco comprometido", "Operar com cautela. Expectativa parcial com o paciente."),
        ("< 0,50", "Tipo 2 — Discordante", "Eixos fragmentados, HOA dominantes ou magnitude dissociada", "Redefinir objetivo. Considerar CXL, lente escleral ou anel tectônico.")
    ]
    
    for r_idx, row_data in enumerate(class_rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl_class.cell(r_idx + 1, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = True

    # 4. FUNDAMENTO BIOMECÂNICO — MODELOS FEM
    add_sec_heading("4", "Fundamento Biomecânico — Modelos FEM")
    add_p("Para validar a premissa física do ICE — de que o desalinhamento entre o eixo do anel e o eixo real do cone reduz a eficácia da correção — foram desenvolvidos 24 modelos de elementos finitos utilizando o software FEBio com material anisotrópico Holzapfel-Gasser-Ogden (HGO), que reproduz a organização fibrilar do estroma.")
    
    add_p("O anel alinhado ao eixo principal do cone gera resposta apical 1,24x mais eficiente do que o anel desalinhado em 30°. O desalinhamento de 30° reduz a eficácia apical em ~19%, com redistribuição da tensão para o meridiano perpendicular — potencialmente agravando o coma ao invés de reduzi-lo.", bold_prefix="Resultado principal — FEM: ")
    add_p("ICE alto significa que o eixo convencional de referência (K-steep) e o eixo real do cone estão alinhados — o anel posicionado pelo nomograma padrão está no campo de força correto. ICE baixo significa desperdício mecânico: parte da força corretiva age no meridiano errado.", bold_prefix="Implicação clínica direta: ")
    
    # Dissociation Note
    tbl_note = doc.add_table(rows=1, cols=1)
    tbl_note.columns[0].width = Inches(6.5)
    cell_note = tbl_note.cell(0, 0)
    set_cell_background(cell_note, "F0FBFB")
    set_cell_left_border(cell_note, "17A1A5", "24")
    set_cell_margins(cell_note, top=100, bottom=100, left=150, right=150)
    p_note = cell_note.paragraphs[0]
    run_nl = p_note.add_run("NOTA SOBRE A DISSOCIAÇÃO ANTERIOR-POSTERIOR:\n")
    run_nl.bold = True
    run_nl.font.size = Pt(9)
    run_nl.font.color.rgb = COLOR_TEAL
    p_note.add_run("Na análise de 660 olhos, a correlação entre posição do ápice anterior e posterior foi de r = 0,028 (p = 0,50 — nula). Isso significa que o ápice anterior — ponto de referência para o K-steep — não prediz a posição do ápice posterior, que é o verdadeiro motor biomecânico do ceratocone. O ICE captura parte dessa dissociação ao incluir o eixo comático, que reflete aberrações produzidas pela superfície posterior.")

    # 5. DADOS PRELIMINARES (RETROSPECTIVO)
    add_sec_heading("5", "Dados Preliminares (Retrospectivo)")
    add_p("A análise retrospectiva utilizou dados de 8 estudos publicados entre 2013 e 2024, totalizando 300 olhos com ICRS, recodificados segundo os critérios ICE a partir dos dados disponibilizados nos artigos. Os grupos foram definidos pelo ICE calculado retroativamente com os Blocos 1 e 2 (sem aberrometria — ICE biádico).")
    
    add_p("Resultados por Grupo de ICE:", space_after=4)
    add_bullet(" +4,2 linhas de BCVA médio.", "ICE Alto (n = 118):")
    add_bullet(" +2,8 linhas de BCVA médio.", "ICE Moderado (n = 102):")
    add_bullet(" +1,6 linhas de BCVA médio.", "ICE Baixo (n = 80):")
    
    # ROC table
    add_p("Desempenho de Predição (Ganho de BCVA >= 3 linhas Snellen):", space_after=4)
    tbl_roc = doc.add_table(rows=5, cols=4)
    tbl_roc.style = 'Light Shading Accent 1'
    
    headers_roc = ["Métrica", "AUC (Área Sob a Curva)", "Intervalo de Confiança", "Comparação com ICE"]
    for i, h in enumerate(headers_roc):
        tbl_roc.cell(0, i).paragraphs[0].add_run(h).bold = True
        
    roc_rows = [
        ("ICE", "0,82", "0,76 – 0,87", "—"),
        ("Kmax", "0,68", "0,61 – 0,75", "p = 0,012 (DeLong)"),
        ("Paquimetria Mínima", "0,64", "0,57 – 0,71", "p = 0,004 (DeLong)"),
        ("Estadiamento Amsler-Krumeich", "0,61", "0,54 – 0,68", "p < 0,001 (DeLong)")
    ]
    
    for r_idx, row_data in enumerate(roc_rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl_roc.cell(r_idx + 1, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = True

    # 6. PROTOCOLO DE ESTUDO PROSPECTIVO PROPOSTO
    add_sec_heading("6", "Protocolo de Estudo Prospectivo Proposto")
    add_p("Delineamento: Estudo de coorte prospectivo, multicêntrico, com cegamento do cirurgião para o resultado do ICE até o desfecho primário. Acompanhamento aos 6 meses pós-operatórios.")
    
    add_p("Critérios de Seleção:", space_after=4)
    add_bullet("Ceratocone confirmado (BAD-D >= 1,6 ou Kmax >= 47,2 D).", "Inclusão: ")
    add_bullet("Candidato a ICRS por indicação clínica independente; CDVA <= 20/40; Idade >= 18 anos.", "Inclusão: ")
    add_bullet("Dados completos disponíveis: Pentacam + aberrometria + refração manifesta.", "Inclusão: ")
    add_bullet("Cirurgia corneana prévia; Paquimetria mínima < 350 µm.", "Exclusão: ")
    add_bullet("Doença ocular associada (glaucoma, uveíte); CXL simultâneo.", "Exclusão: ")
    
    add_p("Dimensionamento Amostral:", space_after=4)
    add_bullet("AUC esperada: 0,82 vs. AUC nula: 0,70. Poder: 80%. Alfa: 0,05.", "Parâmetros: ")
    add_bullet("168 olhos necessários. Proposto com margem de perda: 202 olhos.", "Tamanho: ")
    
    # Outcomes Table
    add_p("Desfechos de Validação Clínicos:", space_after=4)
    tbl_out = doc.add_table(rows=6, cols=3)
    tbl_out.style = 'Light Shading Accent 1'
    
    headers_out = ["Desfecho", "Medida", "Momento"]
    for i, h in enumerate(headers_out):
        tbl_out.cell(0, i).paragraphs[0].add_run(h).bold = True
        
    out_rows = [
        ("Primário", "Ganho de BCVA >= 3 linhas Snellen", "6 meses"),
        ("Secundário 1", "Redução de Kmax >= 2 D", "3 e 6 meses"),
        ("Secundário 2", "Variação de coma RMS", "6 meses"),
        ("Secundário 3", "Satisfação subjetiva (VFQ-25)", "6 meses"),
        ("Secundário 4", "Necessidade de explante ou reposicionamento", "12 meses")
    ]
    
    for r_idx, row_data in enumerate(out_rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl_out.cell(r_idx + 1, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = True

    # 7. CRONOGRAMA PROPOSTO
    add_sec_heading("7", "Cronograma Proposto")
    add_p("Mês 0–2: Cálculo ICE retrospectivo + calibração da fórmula nos 443 pares. Desenvolvimento de rotina automatizada.", bold_prefix="Fase 1 (Preparação): ")
    add_p("Mês 2–4: Protocolo formal, submissão ao CEP do comitê de ética e treinamento das equipes nos centros.", bold_prefix="Fase 2 (Aprovação): ")
    add_p("Mês 4–18: Recrutamento prospectivo e seguimento cego de 202 olhos (avaliações 1m, 3m, 6m).", bold_prefix="Fase 3 (Coleta): ")
    add_p("Mês 18–22: Análise de curvas ROC, regressão multivariada, escrita do artigo e submissão à revista científica (JRS ou Cornea).", bold_prefix="Fase 4 (Escrita/Submissão): ")

    # 8. RECURSOS DESENVOLVIDOS E PUBLICAÇÃO ALVO
    add_sec_heading("8", "Recursos Desenvolvidos e Publicação Alvo")
    add_p("Este projeto visa consolidar a validação do índice ICE através do protocolo prospectivo apresentado. Os recursos e resultados já desenvolvidos estão estruturados para fundamentar uma publicação científica principal.")
    
    tbl_rec = doc.add_table(rows=1, cols=2)
    tbl_rec.style = 'Light Shading Accent 1'
    
    # Left cell - Miguel Reis
    c_left = tbl_rec.cell(0, 0)
    p_left_title = c_left.paragraphs[0]
    p_left_title.add_run("Recursos já Desenvolvidos (Miguel Reis):\n").bold = True
    p_left_title.runs[0].font.size = Pt(9.5)
    p_left_title.runs[0].font.color.rgb = COLOR_TEAL
    
    resources_mr = [
        "Dados retrospectivos de 443 pares de exames.",
        "660 curvaturas SPR analisadas (evidência A-P).",
        "24 modelos biomecânicos FEM anisotrópicos HGO.",
        "Fórmula matemática ICE 2.0 funcional em Python.",
        "Revisão sistemática de 8 estudos clínicos concluída.",
        "Rascunho de artigo inicial estruturado."
    ]
    for r in resources_mr:
        p_r = c_left.add_paragraph(style='List Bullet')
        p_r.paragraph_format.space_after = Pt(2)
        run_r = p_r.add_run(r)
        run_r.font.size = Pt(9)
        
    # Right cell - Technical requirements
    c_right = tbl_rec.cell(0, 1)
    p_right_title = c_right.paragraphs[0]
    p_right_title.add_run("Requisitos para Validação Prospectiva:\n").bold = True
    p_right_title.runs[0].font.size = Pt(9.5)
    p_right_title.runs[0].font.color.rgb = COLOR_TEAL
    
    reqs = [
        "Coorte prospectiva com exames completos (pré e pós).",
        "Dados de aberrometria para cálculo do Bloco 3 (coma).",
        "Acompanhamento de acuidade visual corrigida (6m).",
        "Cegamento estrito da equipe de refração pós-operatória.",
        "Análise estatística ROC final para determinação de limiares.",
        "Ajuste estatístico usando TBI/BAD-D como covariáveis."
    ]
    for rq in reqs:
        p_rq = c_right.add_paragraph(style='List Bullet')
        p_rq.paragraph_format.space_after = Pt(2)
        run_rq = p_rq.add_run(rq)
        run_rq.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # Target Pub Card
    tbl_pub = doc.add_table(rows=1, cols=1)
    tbl_pub.columns[0].width = Inches(6.5)
    cell_pub = tbl_pub.cell(0, 0)
    set_cell_background(cell_pub, "F6F9FC")
    set_cell_margins(cell_pub, top=100, bottom=100, left=150, right=150)
    p_pub = cell_pub.paragraphs[0]
    p_pub.add_run("PUBLICAÇÃO CIENTÍFICA ALVO (ARTIGO PRINCIPAL):\n").bold = True
    p_pub.runs[0].font.size = Pt(9)
    p_pub.runs[0].font.color.rgb = COLOR_NAVY
    p_pub.add_run("• Veículo: Journal of Refractive Surgery ou Cornea\n").font.size = Pt(9.5)
    p_pub.add_run("• Título Proposto: \"Validação prospectiva do Índice de Coerência do Eixo (ICE) como preditor de ganho visual após implante de anel intraestromal em ceratocone\"\n").font.size = Pt(9.5)
    p_pub.add_run("• Escopo: Apresentar a fundamentação biofísica (SPR/FEM), calibração retrospectiva (AUC 0,82) e acurácia diagnóstica prospectiva a 6 meses.").font.size = Pt(9.5)

    # 9. POSIÇÃO HONESTA / LIMITAÇÕES
    add_sec_heading("9", "Posição Honesta — Limitações e Incertezas")
    
    tbl_lim = doc.add_table(rows=7, cols=3)
    tbl_lim.style = 'Light Shading Accent 1'
    
    headers_lim = ["Limitação Identificada", "Status do Projeto", "Estratégia de Mitigação"]
    for i, h in enumerate(headers_lim):
        tbl_lim.cell(0, i).paragraphs[0].add_run(h).bold = True
        
    lim_rows = [
        ("AUC baseada em dados retrospectivos e biádicos", "Parcial", "Validação prospectiva com o índice triádico completo."),
        ("θ_coma ausente na fórmula retrospectiva", "Pendente", "Inclusão de aberrometria sistemática no protocolo prospectivo."),
        ("ICE-score e ICE-min não formalizados de forma independente", "Em revisão", "Artigo metodológico separado definindo ambos matematicamente."),
        ("Limiares de decisão (0,50 / 0,65) derivados de dados externos", "Provisório", "Análise ROC primária do estudo prospectivo redefinirá os pontos."),
        ("Modelos FEM com geometria corneana simplificada", "Robusto", "24 modelos com variação geométrica; resultados concordantes."),
        ("Dissociação A-P da curvatura ainda não publicada", "Pronto", "Artigo metodológico escrito e pronto para submissão.")
    ]
    
    for r_idx, row_data in enumerate(lim_rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl_lim.cell(r_idx + 1, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    add_p("O ICE é, neste momento, uma hipótese biofisicamente fundamentada com suporte retrospectivo moderado. Não é um índice clínico validado. A presente proposta destina-se à avaliação de relevância científica por parte do Prof. Dr. Renato Ambrósio Jr., com o objetivo de colher parecer sobre o potencial de validação do índice no pipeline clínico.", italic=True)

    # Signature block
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    tbl_sig = doc.add_table(rows=1, cols=2)
    
    cell_sig_l = tbl_sig.cell(0, 0)
    p_sig_l = cell_sig_l.paragraphs[0]
    p_sig_l.add_run("Miguel Reis\n").bold = True
    p_sig_l.add_run("Biomecânica Corneana Computacional\nProjeto AVBC-ICRS · Junho 2026\n").font.size = Pt(9)
    p_sig_l.add_run("Investigador Principal").font.size = Pt(8.5)
    p_sig_l.runs[2].italic = True
    
    cell_sig_r = tbl_sig.cell(0, 1)
    p_sig_r = cell_sig_r.paragraphs[0]
    p_sig_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig_r.add_run("Prof. Dr. Renato Ambrósio Jr.\n").bold = True
    p_sig_r.add_run("BrAIn — Rio de Janeiro\nPara Avaliação de Relevância Científica\n").font.size = Pt(9)
    p_sig_r.add_run("Apreciação Científica").font.size = Pt(8.5)
    p_sig_r.runs[2].italic = True
    
    # Save the document
    doc.save(DOC_PATH_LOCAL)
    print(f"Salvo DOCX em: {DOC_PATH_LOCAL}")
    
    # Copy to destination
    shutil.copy2(DOC_PATH_LOCAL, DOC_PATH_DEST)
    print(f"Copiado DOCX para: {DOC_PATH_DEST}")

if __name__ == "__main__":
    build_document()
