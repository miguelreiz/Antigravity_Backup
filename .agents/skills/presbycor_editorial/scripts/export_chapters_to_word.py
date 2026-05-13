#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
"""
PresbyCor Editorial Export Engine
==================================
Equipe Editorial Completa — Padrões Springer Medical / AMA Style
Converte capítulos Markdown para Word (.docx) com qualidade de publicação internacional.

Autor da Skill: Antigravity (Google DeepMind)
Livro: PresbyCor — Dr. Miguel Reis
Versão: 2.0 — Exportação Editorial Rigorosa
"""

import os
import re
import sys
import json
import argparse
import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.oxml
except ImportError:
    print("ERRO: python-docx não instalado. Execute: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# ─────────────────────────────────────────────
# PALETA DE CORES EDITORIAL
# ─────────────────────────────────────────────
# Paleta como tuplas (r,g,b) — mais simples que RGBColor direto
def _rgb(r, g, b):
    """Retorna RGBColor e string hex."""
    return RGBColor(r, g, b), f"{r:02X}{g:02X}{b:02X}"

COLOR_TITLE,    HEX_TITLE    = _rgb(0x1A, 0x3A, 0x5C)
COLOR_HEADING2, HEX_HEADING2 = _rgb(0x1A, 0x3A, 0x5C)
COLOR_HEADING3, HEX_HEADING3 = _rgb(0x2C, 0x5F, 0x8A)
COLOR_HEADING4, HEX_HEADING4 = _rgb(0x2C, 0x5F, 0x8A)
COLOR_TABLE_HDR,HEX_TABLE_HDR= _rgb(0x1A, 0x3A, 0x5C)
COLOR_TABLE_ALT,HEX_TABLE_ALT= _rgb(0xF8, 0xF9, 0xFA)
COLOR_NOTE_BG,  HEX_NOTE_BG  = _rgb(0xEB, 0xF4, 0xFF)
COLOR_NOTE_BDR, HEX_NOTE_BDR = _rgb(0x2C, 0x5F, 0x8A)
COLOR_IMP_BG,   HEX_IMP_BG   = _rgb(0xFF, 0xF8, 0xE1)
COLOR_IMP_BDR,  HEX_IMP_BDR  = _rgb(0xF5, 0x9E, 0x0B)
COLOR_WARN_BG,  HEX_WARN_BG  = _rgb(0xFF, 0xF3, 0xF3)
COLOR_WARN_BDR, HEX_WARN_BDR = _rgb(0xDC, 0x26, 0x26)
COLOR_TIP_BG,   HEX_TIP_BG   = _rgb(0xF0, 0xFD, 0xF4)
COLOR_TIP_BDR,  HEX_TIP_BDR  = _rgb(0x16, 0xA3, 0x4A)


# ─────────────────────────────────────────────
# LOCALIZAÇÃO DO PROJETO
# ─────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).resolve().parent           # scripts/
SKILL_DIR    = SCRIPT_DIR.parent                         # presbycor_editorial/
# .agents/skills/presbycor_editorial/ → go up 3 more to reach Presbycor/
PROJECT_ROOT = SKILL_DIR.parent.parent.parent            # Presbycor/
FIGURES_DIR  = PROJECT_ROOT / "figures"
OUTPUT_DIR   = PROJECT_ROOT / "_Distributable_Book" / "Word_Editorial"
REGISTRY     = SKILL_DIR / "resources" / "chapter_registry.json"

EXPORT_LOG = []


# ═══════════════════════════════════════════════════════════
# HELPERS DE FORMATAÇÃO XML (python-docx de baixo nível)
# ═══════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_str: str):
    """Define cor de fundo de célula de tabela via XML (hex string ex: 'EBF4FF')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_str)
    tcPr.append(shd)


def set_cell_border(cell, side, size=6, color="000000", val="single"):
    """Define borda de célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    border = OxmlElement(f'w:{side}')
    border.set(qn('w:val'), val)
    border.set(qn('w:sz'), str(size))
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), color)
    tcBorders.append(border)


def set_left_border_thick(cell, hex_str: str, size=24):
    """Borda esquerda grossa colorida (para caixas clínicas)."""
    set_cell_border(cell, 'left', size=size, color=hex_str, val='single')
    set_cell_border(cell, 'top', size=4, color='CCCCCC', val='single')
    set_cell_border(cell, 'bottom', size=4, color='CCCCCC', val='single')
    set_cell_border(cell, 'right', size=0, color='FFFFFF', val='none')


def set_run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb


def add_hyperlink_text(para, text):
    """Adiciona texto simples (sem hyperlink real)."""
    run = para.add_run(text)
    return run


# ═══════════════════════════════════════════════════════════
# CONFIGURAÇÃO DO DOCUMENTO BASE
# ═══════════════════════════════════════════════════════════

def create_base_document(chapter_title: str, chapter_num: int) -> Document:
    """Cria documento Word com configurações A4 e estilos editoriais."""
    doc = Document()

    # ── Tamanho A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Cabeçalho
    header = section.header
    header.is_linked_to_previous = False
    htable = header.add_table(1, 2, width=Cm(15.5))
    htable.style = 'Table Grid'
    # Limpar bordas do cabeçalho
    for row in htable.rows:
        for cell in row.cells:
            for side in ['top','bottom','left','right']:
                set_cell_border(cell, side, size=0, color='FFFFFF', val='none')
    hl = htable.cell(0, 0).paragraphs[0]
    hl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hl.add_run("PresbyCor  |  Dr. Miguel Reis")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True

    hr = htable.cell(0, 1).paragraphs[0]
    hr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    from docx.oxml import OxmlElement as OE
    fld = OE('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    instr = OE('w:instrText')
    instr.text = ' PAGE '
    fld_end = OE('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run2 = hr.add_run()
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2._r.append(fld)
    run2._r.append(instr)
    run2._r.append(fld_end)

    # ── Rodapé
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(f"© 2026 Dr. Miguel Reis — Todos os direitos reservados")
    fr.font.name = "Times New Roman"
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fr.font.italic = True

    # ── Estilo Body Text padrão
    style_normal = doc.styles['Normal']
    style_normal.font.name = "Times New Roman"
    style_normal.font.size = Pt(11)
    pf = style_normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15

    return doc


# ═══════════════════════════════════════════════════════════
# APLICAR ESTILOS DE HEADING
# ═══════════════════════════════════════════════════════════

def add_heading(doc: Document, text: str, level: int):
    """Adiciona heading com estilo editorial."""
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    run = para.add_run(text)
    run.font.name = "Times New Roman"

    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_TITLE
        para.paragraph_format.space_before = Pt(24)
        para.paragraph_format.space_after  = Pt(12)
        para.paragraph_format.keep_with_next = True
        # Linha separadora
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1A3A5C')
        pBdr.append(bottom)
        pPr.append(pBdr)

    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_HEADING2
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after  = Pt(8)
        para.paragraph_format.keep_with_next = True

    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = COLOR_HEADING3
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(6)
        para.paragraph_format.keep_with_next = True

    elif level == 4:
        run.font.size = Pt(11)
        run.font.bold = False
        run.font.italic = True
        run.font.color.rgb = COLOR_HEADING4
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after  = Pt(4)
        para.paragraph_format.keep_with_next = True

    return para


# ═══════════════════════════════════════════════════════════
# CAIXAS CLÍNICAS (Alerts)
# ═══════════════════════════════════════════════════════════

ALERT_CONFIG = {
    'NOTE':      ("[NOTA CLINICA]",     HEX_NOTE_BG,  HEX_NOTE_BDR,  COLOR_NOTE_BDR),
    'IMPORTANT': ("[IMPORTANTE]",       HEX_IMP_BG,   HEX_IMP_BDR,   COLOR_IMP_BDR),
    'WARNING':   ("[AVISO]",            HEX_WARN_BG,  HEX_WARN_BDR,  COLOR_WARN_BDR),
    'CAUTION':   ("[ATENCAO]",          HEX_WARN_BG,  HEX_WARN_BDR,  COLOR_WARN_BDR),
    'TIP':       ("[PEROLA CLINICA]",   HEX_TIP_BG,   HEX_TIP_BDR,   COLOR_TIP_BDR),
}

def add_clinical_box(doc: Document, alert_type: str, content_lines: list):
    """Cria caixa clínica editorial com borda colorida."""
    cfg = ALERT_CONFIG.get(alert_type.upper(), ALERT_CONFIG['NOTE'])
    label, hex_bg, hex_bdr, color_bdr = cfg

    # Tabela 1x1 para caixa
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)

    # Borda esquerda grossa colorida
    set_left_border_thick(cell, HEX_NOTE_BDR if alert_type.upper() == 'NOTE' else
                         HEX_IMP_BDR if alert_type.upper() == 'IMPORTANT' else
                         HEX_WARN_BDR if alert_type.upper() in ('WARNING','CAUTION') else
                         HEX_TIP_BDR, size=32)
    set_cell_bg(cell, HEX_NOTE_BG if alert_type.upper() == 'NOTE' else
               HEX_IMP_BG if alert_type.upper() == 'IMPORTANT' else
               HEX_WARN_BG if alert_type.upper() in ('WARNING','CAUTION') else
               HEX_TIP_BG)

    # Conteúdo da célula
    # Primeiro parágrafo: label em negrito
    p_label = cell.paragraphs[0]
    p_label.paragraph_format.space_before = Pt(4)
    p_label.paragraph_format.space_after  = Pt(4)
    run_label = p_label.add_run(label)
    run_label.font.name = "Times New Roman"
    run_label.font.size = Pt(10)
    run_label.font.bold = True
    run_label.font.color.rgb = color_bdr

    # Restante do conteúdo
    text_block = "\n".join(content_lines).strip()
    # Remover marcadores de quote (> )
    text_block = re.sub(r'^>\s*', '', text_block, flags=re.MULTILINE)

    for line in text_block.split('\n'):
        line = line.strip()
        if not line:
            continue
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.3)
        _apply_inline_formatting(p, line, base_size=10.5)

    # Espaço após caixa
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════════════════════
# FORMATAÇÃO INLINE (negrito, itálico, superscript)
# ═══════════════════════════════════════════════════════════

def _apply_inline_formatting(para, text: str, base_size: float = 11, font_name: str = "Times New Roman"):
    """Processa markdown inline: **bold**, *italic*, `code`, [n], superscript."""
    # Regex para tokens de formatação
    pattern = re.compile(
        r'(\*\*\*(?P<boldital>.+?)\*\*\*'
        r'|\*\*(?P<bold>.+?)\*\*'
        r'|\*(?P<ital>.+?)\*'
        r'|`(?P<code>.+?)`'
        r'|\[(?P<refs>[\d,\s\-]+)\]'   # referências [1], [1,2], [1-3]
        r'|(?P<text>[^\[\*`]+)'
        r'|(?P<misc>.+?))',
        re.DOTALL
    )

    for m in pattern.finditer(text):
        g = m.groupdict()
        run = para.add_run()
        run.font.name = font_name
        run.font.size = Pt(base_size)

        if g.get('boldital'):
            run.text = g['boldital']
            run.font.bold = True
            run.font.italic = True
        elif g.get('bold'):
            run.text = g['bold']
            run.font.bold = True
        elif g.get('ital'):
            run.text = g['ital']
            run.font.italic = True
        elif g.get('code'):
            run.text = g['code']
            run.font.name = "Courier New"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif g.get('refs'):
            refs = g['refs']
            run.text = f"[{refs}]"
            run.font.size = Pt(base_size - 2)
            run.font.superscript = True
            run.font.color.rgb = COLOR_HEADING3
        elif g.get('text'):
            run.text = g['text']
        elif g.get('misc'):
            run.text = g['misc']


# ═══════════════════════════════════════════════════════════
# TABELAS MARKDOWN → WORD
# ═══════════════════════════════════════════════════════════

def parse_md_table(lines: list) -> tuple:
    """Extrai cabeçalho e linhas de uma tabela Markdown."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith('|--') or re.match(r'^\|[-\s|:]+\|$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return [], []
    return rows[0], rows[1:]


def add_md_table(doc: Document, md_lines: list, chapter_num: int, table_counter: list):
    """Cria tabela Word formatada a partir de linhas Markdown."""
    header_cells, data_rows = parse_md_table(md_lines)
    if not header_cells:
        return

    table_counter[0] += 1
    n_cols = len(header_cells)
    n_rows = 1 + len(data_rows)

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Largura uniforme
    col_width = Cm(14.5 / n_cols)

    # ── Linha de cabeçalho
    hdr_row = tbl.rows[0]
    for i, cell_text in enumerate(header_cells):
        cell = hdr_row.cells[i]
        cell.width = col_width
        set_cell_bg(cell, HEX_TABLE_HDR)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(cell_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ── Linhas de dados
    for ri, row_data in enumerate(data_rows):
        row = tbl.rows[ri + 1]
        use_alt = (ri % 2 == 1)
        for ci, cell_text in enumerate(row_data):
            if ci >= n_cols:
                break
            cell = row.cells[ci]
            cell.width = col_width
            if use_alt:
                set_cell_bg(cell, HEX_TABLE_ALT)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            _apply_inline_formatting(p, cell_text, base_size=10)

    # Número da tabela
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Tabela {chapter_num}.{table_counter[0]}")
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(9)
    cr.font.bold = True
    cr.font.italic = True
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after  = Pt(12)


# ═══════════════════════════════════════════════════════════
# FIGURAS MARKDOWN → WORD
# ═══════════════════════════════════════════════════════════

def add_figure(doc: Document, alt_text: str, img_path_md: str,
               chapter_num: int, fig_counter: list):
    """Insere figura ou placeholder com legenda editorial."""
    fig_counter[0] += 1
    fig_label = f"Figura {chapter_num}.{fig_counter[0]}"

    # Resolver path da imagem
    img_resolved = None
    if img_path_md:
        # Normalizar path: remover ./
        clean_path = img_path_md.lstrip('./').replace('/', os.sep).replace('\\', os.sep)
        candidate = PROJECT_ROOT / clean_path
        if candidate.exists():
            img_resolved = candidate
        else:
            # Buscar na pasta figures
            fname = Path(img_path_md).name
            for ch_dir in FIGURES_DIR.glob("*"):
                cp = ch_dir / fname
                if cp.exists():
                    img_resolved = cp
                    break

    # Parágrafo de imagem centrado
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(12)
    p_img.paragraph_format.space_after  = Pt(4)

    if img_resolved and HAS_PIL:
        try:
            with Image.open(img_resolved) as im:
                w_px, h_px = im.size
            aspect = h_px / w_px if w_px > 0 else 1.0
            max_w = Cm(14)
            img_h = int(max_w * aspect)
            run = p_img.add_run()
            run.add_picture(str(img_resolved), width=max_w)
            EXPORT_LOG.append(f"  ✓ Figura inserida: {img_resolved.name}")
        except Exception as e:
            _add_figure_placeholder(p_img, fig_label, str(e))
    elif img_resolved:
        try:
            run = p_img.add_run()
            run.add_picture(str(img_resolved), width=Cm(14))
            EXPORT_LOG.append(f"  ✓ Figura inserida: {img_resolved.name}")
        except Exception as e:
            _add_figure_placeholder(p_img, fig_label, str(e))
    else:
        _add_figure_placeholder(p_img, fig_label, "arquivo não encontrado")

    # Legenda
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    run_bold = cap.add_run(f"{fig_label}. ")
    run_bold.font.name = "Times New Roman"
    run_bold.font.size = Pt(9)
    run_bold.font.bold = True
    run_bold.font.italic = True
    run_text = cap.add_run(alt_text[:400] if alt_text else "")  # limitar legenda
    run_text.font.name = "Times New Roman"
    run_text.font.size = Pt(9)
    run_text.font.italic = True


def _add_figure_placeholder(para, label: str, reason: str):
    """Placeholder cinza quando imagem não disponível."""
    run = para.add_run(f"[ {label} — Imagem não disponível: {reason} ]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    EXPORT_LOG.append(f"  ⚠ Placeholder: {label} ({reason})")


# ═══════════════════════════════════════════════════════════
# LISTAS
# ═══════════════════════════════════════════════════════════

def add_list_item(doc: Document, text: str, level: int = 0, ordered: bool = False, num: int = 1):
    """Adiciona item de lista com indentação editorial."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    bullet = f"{num}." if ordered else "•"
    run_bullet = para.add_run(f"{bullet} ")
    run_bullet.font.name = "Times New Roman"
    run_bullet.font.size = Pt(11)
    _apply_inline_formatting(para, text.lstrip('- *•').strip())


# ═══════════════════════════════════════════════════════════
# FÓRMULAS
# ═══════════════════════════════════════════════════════════

def add_formula(doc: Document, formula: str):
    """Adiciona fórmula matemática centrada e formatada."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(8)
    # Limpar LaTeX básico
    clean = formula.strip().strip('$').strip()
    clean = re.sub(r'\\text\{([^}]+)\}', r'\1', clean)
    clean = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', clean)
    clean = re.sub(r'\\times', '×', clean)
    clean = re.sub(r'\\leq', '≤', clean)
    clean = re.sub(r'\\geq', '≥', clean)
    clean = re.sub(r'\\approx', '≈', clean)
    clean = re.sub(r'\\Delta', 'Δ', clean)
    clean = re.sub(r'\\sqrt\{([^}]+)\}', r'√(\1)', clean)
    clean = re.sub(r'\\mu m', 'μm', clean)
    clean = re.sub(r'\\_', '_', clean)
    clean = re.sub(r'\\\^', '^', clean)
    run = para.add_run(clean)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.italic = True


# ═══════════════════════════════════════════════════════════
# SEÇÃO DE REFERÊNCIAS
# ═══════════════════════════════════════════════════════════

def add_references_section(doc: Document, ref_lines: list):
    """Adiciona seção de referências formatada AMA."""
    if not ref_lines:
        return
    # Separador
    doc.add_paragraph()
    add_heading(doc, "Referências Bibliográficas", 2)
    for line in ref_lines:
        line = line.strip()
        if not line:
            continue
        # Remover marcadores e já é AMA
        line = re.sub(r'^\d+\.\s*', '', line)
        # Detectar número
        m = re.match(r'^(\d+)\.\s*(.*)', line.strip())
        if m:
            num, rest = m.group(1), m.group(2)
        else:
            num, rest = None, line

        para = doc.add_paragraph()
        para.paragraph_format.left_indent   = Cm(0.8)
        para.paragraph_format.first_line_indent = Cm(-0.8)
        para.paragraph_format.space_before  = Pt(2)
        para.paragraph_format.space_after   = Pt(3)

        if num:
            r_num = para.add_run(f"{num}. ")
            r_num.font.name = "Times New Roman"
            r_num.font.size = Pt(10)
            r_num.font.bold = True

        # Aplicar itálico nos títulos de journals
        # Detectar padrão: *Journal Name* → itálico
        parts = re.split(r'\*([^*]+)\*', rest)
        for i, part in enumerate(parts):
            r = para.add_run(part)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
            if i % 2 == 1:  # dentro de *...*
                r.font.italic = True


# ═══════════════════════════════════════════════════════════
# PARSER PRINCIPAL DO MARKDOWN
# ═══════════════════════════════════════════════════════════

def parse_and_render(doc: Document, md_text: str, chapter_num: int):
    """
    Parser completo do Markdown do PresbyCor.
    Renderiza todas as construções especiais no documento Word.
    """
    lines = md_text.splitlines()
    n = len(lines)
    i = 0

    fig_counter   = [0]
    table_counter = [0]
    ref_lines     = []
    in_refs       = False
    in_alert      = False
    alert_type    = None
    alert_content = []
    bullet_num    = [0]

    while i < n:
        raw = lines[i]
        line = raw.rstrip()

        # ── SEÇÃO DE REFERÊNCIAS ──────────────────────────────
        if re.match(r'^#+\s*(References|Referências)', line.strip(), re.IGNORECASE):
            in_refs = True
            i += 1
            continue

        if in_refs:
            if re.match(r'^#+\s', line):
                in_refs = False  # nova seção encerra refs
                add_references_section(doc, ref_lines)
                ref_lines = []
                # continua para processar o heading abaixo
            else:
                ref_lines.append(line)
                i += 1
                continue

        # ── ALERT BOX (GitHub style) ──────────────────────────
        alert_start = re.match(r'^>\s*\[!(NOTE|IMPORTANT|WARNING|CAUTION|TIP)\]', line)
        if alert_start and not in_alert:
            in_alert = True
            alert_type = alert_start.group(1)
            alert_content = []
            i += 1
            continue

        if in_alert:
            if line.startswith('>'):
                alert_content.append(line[1:].strip())
                i += 1
                continue
            else:
                # Fim do alert
                add_clinical_box(doc, alert_type, alert_content)
                in_alert = False
                alert_type = None
                alert_content = []
                # não avança i, reprocessa linha atual

        # ── HEADING ──────────────────────────────────────────
        hm = re.match(r'^(#{1,4})\s+(.*)', line)
        if hm:
            level = len(hm.group(1))
            title = hm.group(2).strip()
            # Limpar markdown inline do título
            title = re.sub(r'\*+([^*]+)\*+', r'\1', title)
            add_heading(doc, title, level)
            bullet_num[0] = 0
            i += 1
            continue

        # ── TABELA MARKDOWN ───────────────────────────────────
        if re.match(r'^\s*\|', line) and i < n - 1:
            table_lines = []
            while i < n and re.match(r'^\s*\|', lines[i]):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                add_md_table(doc, table_lines, chapter_num, table_counter)
            continue

        # ── FIGURA ───────────────────────────────────────────
        fig_m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
        if fig_m:
            alt   = fig_m.group(1)
            fpath = fig_m.group(2)
            # Linha seguinte pode ser legenda *...*
            caption_extra = ""
            if i + 1 < n and re.match(r'^\*', lines[i+1].strip()):
                caption_extra = lines[i+1].strip().strip('*').strip()
                i += 1
            full_alt = alt or caption_extra
            add_figure(doc, full_alt, fpath, chapter_num, fig_counter)
            i += 1
            continue

        # Legenda de figura isolada (linha começa com *)
        if re.match(r'^\*Figura\s', line) or re.match(r'^\*Figure\s', line):
            # já tratado acima, ignorar
            i += 1
            continue

        # ── FÓRMULA LaTeX ─────────────────────────────────────
        if line.strip().startswith('$$') and line.strip().endswith('$$') and len(line.strip()) > 4:
            formula = line.strip()[2:-2]
            add_formula(doc, formula)
            i += 1
            continue

        if line.strip() == '$$' or line.strip().startswith('$$'):
            formula_lines = []
            if line.strip().startswith('$$') and not line.strip().endswith('$$'):
                i += 1
                while i < n and '$$' not in lines[i]:
                    formula_lines.append(lines[i].strip())
                    i += 1
            formula = ' '.join(formula_lines)
            if formula:
                add_formula(doc, formula)
            i += 1
            continue

        # ── SEPARADOR HORIZONTAL ──────────────────────────────
        if re.match(r'^---+$', line.strip()) or re.match(r'^___+$', line.strip()):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after  = Pt(6)
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── LISTA NÃO-ORDENADA ────────────────────────────────
        ul_m = re.match(r'^(\s*)([-*+])\s+(.*)', line)
        if ul_m:
            indent = len(ul_m.group(1)) // 2
            content = ul_m.group(3)
            add_list_item(doc, content, level=indent, ordered=False)
            bullet_num[0] = 0
            i += 1
            continue

        # ── LISTA ORDENADA ────────────────────────────────────
        ol_m = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if ol_m:
            indent = len(ol_m.group(1)) // 2
            num    = int(ol_m.group(2))
            content = ol_m.group(3)
            add_list_item(doc, content, level=indent, ordered=True, num=num)
            i += 1
            continue

        # ── LINHA EM BRANCO ───────────────────────────────────
        if not line.strip():
            # Não adicionar parágrafos em branco excessivos
            i += 1
            continue

        # ── PARÁGRAFO NORMAL ──────────────────────────────────
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(6)
        para.paragraph_format.first_line_indent = Cm(0.5)
        _apply_inline_formatting(para, line)
        i += 1

    # Fechar alert pendente
    if in_alert and alert_content:
        add_clinical_box(doc, alert_type or 'NOTE', alert_content)

    # Referências pendentes
    if ref_lines:
        add_references_section(doc, ref_lines)


# ═══════════════════════════════════════════════════════════
# EXPORTAR UM CAPÍTULO
# ═══════════════════════════════════════════════════════════

def export_chapter(chapter_info: dict, verbose: bool = False) -> bool:
    """Exporta um capítulo Markdown para Word editorial."""
    chap_id   = chapter_info['id']
    chap_file = PROJECT_ROOT / chapter_info['file']
    out_name  = chapter_info['output']
    title     = chapter_info['title']
    out_path  = OUTPUT_DIR / out_name

    print(f"\n🔷 Exportando: {out_name}")
    EXPORT_LOG.append(f"\n{'='*60}")
    EXPORT_LOG.append(f"CAPÍTULO {chap_id}: {title}")
    EXPORT_LOG.append(f"Arquivo fonte: {chap_file.name}")
    EXPORT_LOG.append(f"Saída: {out_name}")

    if not chap_file.exists():
        msg = f"  ✗ Arquivo não encontrado: {chap_file}"
        print(msg)
        EXPORT_LOG.append(msg)
        return False

    # Ler conteúdo
    with open(chap_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Criar documento
    doc = create_base_document(title, chap_id)

    # Título do livro na primeira página
    if chap_id == 0:
        p_book = doc.add_paragraph()
        p_book.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_book.add_run("PRESBYCOR")
        r.font.name = "Times New Roman"
        r.font.size = Pt(24)
        r.font.bold = True
        r.font.color.rgb = COLOR_TITLE
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("Modern Strategies for Presbyopia and Laser Mechanics")
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(14)
        r2.font.italic = True
        r2.font.color.rgb = COLOR_HEADING3
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run("Dr. Miguel Reis")
        r3.font.name = "Times New Roman"
        r3.font.size = Pt(12)
        r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        doc.add_paragraph()

    # Processar Markdown
    parse_and_render(doc, md_content, chap_id)

    # Salvar
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    size_kb = out_path.stat().st_size // 1024
    msg = f"  ✓ Salvo: {out_path.name} ({size_kb} KB)"
    print(msg)
    EXPORT_LOG.append(msg)
    return True


# ═══════════════════════════════════════════════════════════
# RELATÓRIO DE EXPORTAÇÃO
# ═══════════════════════════════════════════════════════════

def save_report():
    """Salva relatório completo de exportação."""
    report_path = OUTPUT_DIR / "EXPORT_REPORT.txt"
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    header = [
        "=" * 60,
        "PRESBYCOR — RELATÓRIO DE EXPORTAÇÃO EDITORIAL",
        f"Data: {now}",
        f"Engine: export_chapters_to_word.py v2.0",
        "=" * 60,
    ]
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(header + EXPORT_LOG))
    print(f"\n📄 Relatório salvo: {report_path}")


# ═══════════════════════════════════════════════════════════
# ENTRY POINT
# ═══════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="PresbyCor Editorial Export — Word Publishing Engine v2.0"
    )
    parser.add_argument('--all', action='store_true',
                        help='Exportar todos os capítulos')
    parser.add_argument('--chapter', type=int, metavar='N',
                        help='Exportar capítulo específico (0=Prefácio, 1-13=Capítulos)')
    parser.add_argument('--verbose', action='store_true',
                        help='Log detalhado')
    args = parser.parse_args()

    # Carregar registry
    if not REGISTRY.exists():
        print(f"ERRO: Registry não encontrado em {REGISTRY}")
        sys.exit(1)

    with open(REGISTRY, 'r', encoding='utf-8') as f:
        registry = json.load(f)

    chapters = registry['chapters']

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"\n[PresbyCor] Editorial Export Engine v2.0")
    print(f"[Projeto] {PROJECT_ROOT}")
    print(f"[Saida]   {OUTPUT_DIR}")

    success = 0
    failed  = 0

    if args.all:
        print(f"\n>> Exportando {len(chapters)} capitulos...\n")
        for ch in chapters:
            ok = export_chapter(ch, verbose=args.verbose)
            if ok:
                success += 1
            else:
                failed += 1

    elif args.chapter is not None:
        chap = next((c for c in chapters if c['id'] == args.chapter), None)
        if not chap:
            print(f"ERRO: Capítulo {args.chapter} não encontrado no registry.")
            sys.exit(1)
        ok = export_chapter(chap, verbose=args.verbose)
        success = 1 if ok else 0
        failed  = 0 if ok else 1

    else:
        parser.print_help()
        sys.exit(0)

    # Resumo
    print(f"\n{'='*50}")
    print(f"[OK] Exportados com sucesso: {success}")
    if failed:
        print(f"[ERRO] Com erros: {failed}")
    print(f"[DIR] Arquivos em: {OUTPUT_DIR}")

    EXPORT_LOG.append(f"\n{'='*60}")
    EXPORT_LOG.append(f"RESUMO: {success} exportados, {failed} com erro")

    save_report()


if __name__ == '__main__':
    main()
