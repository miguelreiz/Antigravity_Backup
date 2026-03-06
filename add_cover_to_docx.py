#!/usr/bin/env python3
"""
Adiciona capa ao DOCX do livro PresbyCor
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_cover_to_docx(docx_path, cover_image_path, output_path):
    """
    Adiciona uma capa de imagem ao início de um documento DOCX existente
    """
    print(f"📖 Abrindo documento: {docx_path}")
    doc = Document(docx_path)
    
    # Criar novo parágrafo no início
    print("🖼️  Inserindo capa...")
    first_paragraph = doc.paragraphs[0]._element
    first_paragraph.addprevious(doc.add_paragraph()._element)
    
    # Adicionar imagem da capa
    paragraph = doc.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(cover_image_path, width=Inches(6.5))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar quebra de página após a capa
    paragraph.add_run().add_break()
    
    # Salvar
    print(f"💾 Salvando como: {output_path}")
    doc.save(output_path)
    print("✅ Concluído!")

if __name__ == "__main__":
    import sys
    import os
    
    # Caminhos
    base_dir = "/Users/miguelreis/Downloads/Takeout/NotebookLM/PresbyCor_ Modern Strategies for Presbyopia and La"
    docx_path = os.path.join(base_dir, "_Export_To_Drive/PresbyCor_MASTER_BOOK.docx")
    cover_path = os.path.join(base_dir, "presbiopia_laser_cover_pt.png")
    output_path = "/Users/miguelreis/Desktop/PresbyCor_v1.0_COMPLETO_COM_CAPA.docx"
    
    # Verificar arquivos
    if not os.path.exists(docx_path):
        print(f"❌ DOCX não encontrado: {docx_path}")
        sys.exit(1)
    
    if not os.path.exists(cover_path):
        print(f"❌ Capa não encontrada: {cover_path}")
        sys.exit(1)
    
    # Processar
    add_cover_to_docx(docx_path, cover_path, output_path)
    
    # Mostrar tamanho final
    size_mb = os.path.getsize(output_path) / (1024 * 1024)
    print(f"\n📊 Arquivo final: {size_mb:.1f} MB")
    print(f"📁 Localização: {output_path}")
